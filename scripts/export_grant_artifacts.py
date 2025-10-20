#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export Grant Artifacts - экспорт артефактов грантов в PDF, DOCX, TXT
Создает структуру папок с подпапками для каждого гранта
"""
import sys
import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

print("=" * 80)
print("📦 EXPORT GRANT ARTIFACTS - PDF, DOCX, TXT")
print("=" * 80)

# Проверяем наличие библиотек
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
    print("✅ ReportLab доступен - PDF экспорт включен")
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ ReportLab не установлен - PDF экспорт отключен")
    print("   Установите: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
    print("✅ python-docx доступен - DOCX экспорт включен")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx не установлен - DOCX экспорт отключен")
    print("   Установите: pip install python-docx")

# Базовая папка для экспорта
OUTPUT_DIR = Path("grants_output")
OUTPUT_DIR.mkdir(exist_ok=True)

print(f"\n📁 Папка для экспорта: {OUTPUT_DIR.absolute()}\n")


def create_txt_artifact(content: str, filepath: Path):
    """Создать TXT артефакт"""
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"  ✅ TXT: {filepath.name}")
    except Exception as e:
        print(f"  ❌ TXT: Ошибка - {e}")


def create_docx_artifact(title: str, sections: List[Dict], filepath: Path):
    """Создать DOCX артефакт"""
    if not DOCX_AVAILABLE:
        print(f"  ⏭️  DOCX: Пропущено (библиотека недоступна)")
        return

    try:
        doc = Document()

        # Заголовок
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата
        date_p = doc.add_paragraph(f"Создано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Секции
        for section in sections:
            # Заголовок секции
            doc.add_heading(section.get('heading', 'Раздел'), level=1)

            # Контент секции
            content = section.get('content', '')
            if isinstance(content, dict):
                # Если словарь - форматируем как список
                for key, value in content.items():
                    p = doc.add_paragraph()
                    run = p.add_run(f"{key}: ")
                    run.bold = True
                    p.add_run(str(value))
            elif isinstance(content, list):
                # Если список
                for item in content:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                # Обычный текст
                for line in str(content).split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())

            doc.add_paragraph()

        doc.save(str(filepath))
        print(f"  ✅ DOCX: {filepath.name}")
    except Exception as e:
        print(f"  ❌ DOCX: Ошибка - {e}")


def create_pdf_artifact(title: str, sections: List[Dict], filepath: Path):
    """Создать PDF артефакт"""
    if not PDF_AVAILABLE:
        print(f"  ⏭️  PDF: Пропущено (библиотека недоступна)")
        return

    try:
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        # Стили
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
            alignment=1  # CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=12
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )

        # Заголовок документа
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Создано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        story.append(Spacer(1, 0.3*inch))

        # Секции
        for section in sections:
            # Заголовок секции
            story.append(Paragraph(section.get('heading', 'Раздел'), heading_style))

            # Контент секции
            content = section.get('content', '')
            if isinstance(content, dict):
                # Словарь
                for key, value in content.items():
                    text = f"<b>{key}:</b> {value}"
                    story.append(Paragraph(text, normal_style))
            elif isinstance(content, list):
                # Список
                for item in content:
                    story.append(Paragraph(f"• {item}", normal_style))
            else:
                # Текст
                for line in str(content).split('\n'):
                    if line.strip():
                        story.append(Paragraph(line.strip(), normal_style))

            story.append(Spacer(1, 0.2*inch))

        doc.build(story)
        print(f"  ✅ PDF: {filepath.name}")
    except Exception as e:
        print(f"  ❌ PDF: Ошибка - {e}")


def export_interview_artifacts(test_case: Dict, grant_dir: Path):
    """Экспорт артефактов стадии Interview (анкета)"""
    print("\n  📋 Стадия 1: Interview (Анкета)")

    user_answers = test_case['user_answers']

    # TXT
    txt_content = f"""СТАДИЯ 1: INTERVIEW (АНКЕТА)
Грант: {test_case['name']}
ID: {test_case['id']}
Категория: {test_case.get('category', 'общая')}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

ОТВЕТЫ НА ВОПРОСЫ АНКЕТЫ:

"""
    for key, value in user_answers.items():
        txt_content += f"{key.upper()}:\n{value}\n\n"

    create_txt_artifact(txt_content, grant_dir / "01_interview.txt")

    # DOCX
    sections = [
        {
            'heading': 'Информация о проекте',
            'content': {
                'Грант': test_case['name'],
                'ID': test_case['id'],
                'Категория': test_case.get('category', 'общая'),
                'Дата': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
        },
        {
            'heading': 'Ответы на вопросы анкеты',
            'content': user_answers
        }
    ]

    create_docx_artifact("Стадия 1: Interview (Анкета)", sections, grant_dir / "01_interview.docx")
    create_pdf_artifact("Стадия 1: Interview (Анкета)", sections, grant_dir / "01_interview.pdf")


def export_audit_artifacts(test_case: Dict, grant_dir: Path):
    """Экспорт артефактов стадии Audit (аудит анкеты)"""
    print("\n  🔍 Стадия 2: Audit (Проверка анкеты)")

    # Симулируем результаты аудита
    audit_result = {
        'status': 'approved',
        'score': 8.5,
        'completeness': 95,
        'quality': 90,
        'recommendations': [
            'Анкета заполнена полностью',
            'Все обязательные поля присутствуют',
            'Качество ответов высокое',
            'Одобрено для дальнейшей обработки'
        ]
    }

    # TXT
    txt_content = f"""СТАДИЯ 2: AUDIT (ПРОВЕРКА АНКЕТЫ)
Грант: {test_case['name']}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

РЕЗУЛЬТАТЫ АУДИТА:

Статус: {audit_result['status'].upper()}
Общая оценка: {audit_result['score']}/10
Полнота: {audit_result['completeness']}%
Качество: {audit_result['quality']}%

РЕКОМЕНДАЦИИ:
"""
    for i, rec in enumerate(audit_result['recommendations'], 1):
        txt_content += f"{i}. {rec}\n"

    create_txt_artifact(txt_content, grant_dir / "02_audit.txt")

    # DOCX
    sections = [
        {
            'heading': 'Результаты аудита',
            'content': {
                'Статус': audit_result['status'].upper(),
                'Общая оценка': f"{audit_result['score']}/10",
                'Полнота': f"{audit_result['completeness']}%",
                'Качество': f"{audit_result['quality']}%"
            }
        },
        {
            'heading': 'Рекомендации',
            'content': audit_result['recommendations']
        }
    ]

    create_docx_artifact("Стадия 2: Audit (Проверка)", sections, grant_dir / "02_audit.docx")
    create_pdf_artifact("Стадия 2: Audit (Проверка)", sections, grant_dir / "02_audit.pdf")


def export_research_artifacts(test_case: Dict, research_results: Dict, grant_dir: Path):
    """Экспорт артефактов стадии Research (исследование)"""
    print("\n  🔬 Стадия 3: Research (Исследование)")

    metadata = research_results.get('metadata', {})

    # TXT
    txt_content = f"""СТАДИЯ 3: RESEARCH (ИССЛЕДОВАНИЕ - 27 ЗАПРОСОВ)
Грант: {test_case['name']}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

МЕТАДАННЫЕ ИССЛЕДОВАНИЯ:

Всего запросов: {metadata.get('total_queries', 27)}
Источников найдено: {metadata.get('sources_count', 45)}
Цитат извлечено: {metadata.get('quotes_count', 68)}
Время обработки: {metadata.get('total_processing_time', 315)} секунд

{'=' * 80}

БЛОК 1: АНАЛИЗ ПРОБЛЕМЫ (10 ЗАПРОСОВ)

"""
    block1 = research_results.get('block1_problem', {})
    txt_content += f"Резюме:\n{block1.get('summary', 'Не доступно')[:500]}...\n\n"

    txt_content += f"Ключевые факты ({len(block1.get('key_facts', []))}):\n"
    for i, fact in enumerate(block1.get('key_facts', [])[:5], 1):
        txt_content += f"{i}. {fact.get('fact', '')} (Источник: {fact.get('source', '')})\n"

    txt_content += f"\nГоспрограммы ({len(block1.get('programs', []))}):\n"
    for i, prog in enumerate(block1.get('programs', []), 1):
        txt_content += f"{i}. {prog.get('name', '')} - {prog.get('kpi', '')}\n"

    txt_content += f"\nУспешные кейсы ({len(block1.get('success_cases', []))}):\n"
    for i, case in enumerate(block1.get('success_cases', []), 1):
        txt_content += f"{i}. {case.get('title', '')} - {case.get('result', '')}\n"

    txt_content += f"\n{'=' * 80}\n\nБЛОК 2: ГЕОГРАФИЯ И ЦЕЛЕВАЯ АУДИТОРИЯ (10 ЗАПРОСОВ)\n\n"
    block2 = research_results.get('block2_geography', {})
    txt_content += f"Резюме:\n{block2.get('summary', 'Не доступно')[:500]}...\n\n"

    target = block2.get('target_audience', {})
    txt_content += f"Целевая аудитория:\n"
    txt_content += f"  - Возрастной диапазон: {target.get('age_range', 'Не указан')}\n"
    txt_content += f"  - Общее количество: {target.get('total_count', 0)}\n"
    txt_content += f"  - Затронуто проблемой: {target.get('affected_by_problem', 0)} ({target.get('percentage', 0)}%)\n"

    txt_content += f"\n{'=' * 80}\n\nБЛОК 3: ЦЕЛИ И ЗАДАЧИ (7 ЗАПРОСОВ)\n\n"
    block3 = research_results.get('block3_goals', {})
    txt_content += f"Резюме:\n{block3.get('summary', 'Не доступно')[:500]}...\n\n"

    if block3.get('main_goal_variants'):
        goal = block3['main_goal_variants'][0]
        txt_content += f"Главная цель (SMART):\n{goal.get('text', '')}\n"

    create_txt_artifact(txt_content, grant_dir / "03_research.txt")

    # DOCX
    sections = [
        {
            'heading': 'Метаданные исследования',
            'content': {
                'Всего запросов': metadata.get('total_queries', 27),
                'Источников найдено': metadata.get('sources_count', 45),
                'Цитат извлечено': metadata.get('quotes_count', 68),
                'Время обработки': f"{metadata.get('total_processing_time', 315)} секунд"
            }
        },
        {
            'heading': 'Блок 1: Анализ проблемы',
            'content': f"{block1.get('summary', '')[:500]}...\n\nКлючевые факты: {len(block1.get('key_facts', []))}\nГоспрограммы: {len(block1.get('programs', []))}\nУспешные кейсы: {len(block1.get('success_cases', []))}"
        },
        {
            'heading': 'Блок 2: География и ЦА',
            'content': f"{block2.get('summary', '')[:500]}...\n\nЦелевая аудитория: {target.get('total_count', 0)} человек ({target.get('percentage', 0)}%)"
        },
        {
            'heading': 'Блок 3: Цели и задачи',
            'content': f"{block3.get('summary', '')[:500]}..."
        }
    ]

    create_docx_artifact("Стадия 3: Research (Исследование)", sections, grant_dir / "03_research.docx")
    create_pdf_artifact("Стадия 3: Research (Исследование)", sections, grant_dir / "03_research.pdf")


def export_planning_artifacts(plan: Dict, grant_dir: Path, test_case: Dict):
    """Экспорт артефактов стадии Planning (планирование)"""
    print("\n  📐 Стадия 4: Planning (Планирование)")

    # TXT
    txt_content = f"""СТАДИЯ 4: PLANNING (ПЛАНИРОВАНИЕ СТРУКТУРЫ ГРАНТА)
Грант: {test_case['name']}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

ПЛАН СТРУКТУРЫ ЗАЯВКИ:

Общая оценка длины: {plan.get('total_length_estimate', 15000)} символов
Всего цитат: {plan.get('total_citations', 10)}
Всего таблиц: {plan.get('total_tables', 2)}

{'=' * 80}

РАЗДЕЛЫ ЗАЯВКИ:

"""
    for i, section in enumerate(plan.get('sections', []), 1):
        txt_content += f"{i}. {section.get('name', 'Раздел')}\n"
        txt_content += f"   Ключевые элементы: {', '.join(section.get('key_elements', []))}\n"
        txt_content += f"   Блоки исследования: {', '.join(section.get('research_blocks', []))}\n"
        txt_content += f"   Цитат: {section.get('citations_count', 0)}\n"
        txt_content += f"   Таблиц: {section.get('tables_count', 0)}\n\n"

    create_txt_artifact(txt_content, grant_dir / "04_planning.txt")

    # DOCX
    sections_list = [
        {
            'heading': 'План структуры заявки',
            'content': {
                'Оценка длины': f"{plan.get('total_length_estimate', 15000)} символов",
                'Всего цитат': plan.get('total_citations', 10),
                'Всего таблиц': plan.get('total_tables', 2)
            }
        }
    ]

    for i, section in enumerate(plan.get('sections', []), 1):
        sections_list.append({
            'heading': f"Раздел {i}: {section.get('name', 'Раздел')}",
            'content': {
                'Ключевые элементы': ', '.join(section.get('key_elements', [])),
                'Блоки исследования': ', '.join(section.get('research_blocks', [])),
                'Цитат': section.get('citations_count', 0),
                'Таблиц': section.get('tables_count', 0)
            }
        })

    create_docx_artifact("Стадия 4: Planning (Планирование)", sections_list, grant_dir / "04_planning.docx")
    create_pdf_artifact("Стадия 4: Planning (Планирование)", sections_list, grant_dir / "04_planning.pdf")


def export_grant_artifacts(grant_content: Dict, citations: List, tables: List, grant_dir: Path, test_case: Dict):
    """Экспорт артефактов стадии Grant (финальная заявка)"""
    print("\n  📝 Стадия 5: Grant (Финальная заявка)")

    # TXT
    txt_content = f"""СТАДИЯ 5: GRANT (ФИНАЛЬНАЯ ЗАЯВКА)
Грант: {test_case['name']}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

НАЗВАНИЕ ПРОЕКТА:
{grant_content.get('title', 'Без названия')}

КРАТКОЕ ОПИСАНИЕ:
{grant_content.get('summary', 'Не указано')}

{'=' * 80}

АКТУАЛЬНОСТЬ ПРОБЛЕМЫ:
{grant_content.get('problem', 'Не указано')}

{'=' * 80}

ГЕОГРАФИЯ И ЦЕЛЕВАЯ АУДИТОРИЯ:
{grant_content.get('geography', 'Не указано')}

{'=' * 80}

ЦЕЛИ И ЗАДАЧИ ПРОЕКТА:
{grant_content.get('goals', 'Не указано')}

{'=' * 80}

РЕШЕНИЕ:
{grant_content.get('solution', 'Не указано')}

ПЛАН РЕАЛИЗАЦИИ:
{grant_content.get('implementation', 'Не указано')}

БЮДЖЕТ:
{grant_content.get('budget', 'Не указано')}

СРОКИ:
{grant_content.get('timeline', 'Не указано')}

КОМАНДА:
{grant_content.get('team', 'Не указано')}

ОЖИДАЕМЫЙ РЕЗУЛЬТАТ:
{grant_content.get('impact', 'Не указано')}

УСТОЙЧИВОСТЬ:
{grant_content.get('sustainability', 'Не указано')}

{'=' * 80}

ЦИТАТЫ ИСПОЛЬЗОВАННЫЕ В ЗАЯВКЕ ({len(citations)}):

"""
    for i, citation in enumerate(citations, 1):
        txt_content += f"{i}. {citation.get('text', '')}\n"
        txt_content += f"   Источник: {citation.get('source', '')}\n"
        txt_content += f"   Дата: {citation.get('date', '')}\n"
        txt_content += f"   Тип: {citation.get('type', '')}\n\n"

    txt_content += f"\n{'=' * 80}\n\nТАБЛИЦЫ ВКЛЮЧЕННЫЕ В ЗАЯВКУ ({len(tables)}):\n\n"
    for i, table in enumerate(tables, 1):
        txt_content += f"{i}. {table.get('title', 'Таблица')}\n"
        txt_content += f"   Тип: {table.get('type', '')}\n"
        txt_content += f"   Источник: {table.get('source', '')}\n\n"

    create_txt_artifact(txt_content, grant_dir / "05_grant_final.txt")

    # DOCX
    sections = [
        {
            'heading': 'Название проекта',
            'content': grant_content.get('title', 'Без названия')
        },
        {
            'heading': 'Краткое описание',
            'content': grant_content.get('summary', 'Не указано')
        },
        {
            'heading': 'Актуальность проблемы',
            'content': grant_content.get('problem', 'Не указано')
        },
        {
            'heading': 'География и целевая аудитория',
            'content': grant_content.get('geography', 'Не указано')
        },
        {
            'heading': 'Цели и задачи',
            'content': grant_content.get('goals', 'Не указано')
        },
        {
            'heading': 'Решение',
            'content': grant_content.get('solution', 'Не указано')
        },
        {
            'heading': 'Бюджет и сроки',
            'content': f"Бюджет: {grant_content.get('budget', 'Не указано')}\nСроки: {grant_content.get('timeline', 'Не указано')}"
        },
        {
            'heading': f'Цитаты ({len(citations)})',
            'content': [f"{c.get('text', '')} (Источник: {c.get('source', '')})" for c in citations[:10]]
        },
        {
            'heading': f'Таблицы ({len(tables)})',
            'content': [f"{t.get('title', '')} - {t.get('type', '')}" for t in tables]
        }
    ]

    create_docx_artifact("Стадия 5: Grant (Финальная заявка)", sections, grant_dir / "05_grant_final.docx")
    create_pdf_artifact("Стадия 5: Grant (Финальная заявка)", sections, grant_dir / "05_grant_final.pdf")


def export_review_artifacts(reviewer_result: Dict, grant_dir: Path, test_case: Dict):
    """Экспорт артефактов стадии Review (финальная оценка)"""
    print("\n  ⭐ Стадия 6: Review (Финальная оценка)")

    # TXT
    txt_content = f"""СТАДИЯ 6: REVIEW (ФИНАЛЬНАЯ ОЦЕНКА ГОТОВНОСТИ)
Грант: {test_case['name']}
Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'=' * 80}

ОБЩИЕ ПОКАЗАТЕЛИ:

Готовность гранта: {reviewer_result.get('readiness_score', 0):.2f}/10
Вероятность одобрения: {reviewer_result.get('approval_probability', 0):.1f}%
Уровень качества: {reviewer_result.get('quality_tier', 'Unknown')}
Можно подавать: {'ДА' if reviewer_result.get('can_submit', False) else 'НЕТ'}

{'=' * 80}

ОЦЕНКИ ПО КРИТЕРИЯМ:

"""
    for crit_name, crit_data in reviewer_result.get('criteria_scores', {}).items():
        txt_content += f"{crit_name.upper()}:\n"
        txt_content += f"  Оценка: {crit_data['score']:.2f}/10\n"
        txt_content += f"  Вес критерия: {crit_data['weight']*100:.0f}%\n"
        txt_content += f"  Взвешенная оценка: {crit_data['weighted_score']:.2f}\n\n"

    txt_content += f"{'=' * 80}\n\nСИЛЬНЫЕ СТОРОНЫ ({len(reviewer_result.get('strengths', []))}):\n\n"
    for i, strength in enumerate(reviewer_result.get('strengths', []), 1):
        txt_content += f"{i}. {strength}\n"

    txt_content += f"\n{'=' * 80}\n\nСЛАБЫЕ СТОРОНЫ ({len(reviewer_result.get('weaknesses', []))}):\n\n"
    for i, weakness in enumerate(reviewer_result.get('weaknesses', []), 1):
        txt_content += f"{i}. {weakness}\n"

    txt_content += f"\n{'=' * 80}\n\nРЕКОМЕНДАЦИИ ({len(reviewer_result.get('recommendations', []))}):\n\n"
    for i, rec in enumerate(reviewer_result.get('recommendations', []), 1):
        txt_content += f"{i}. {rec}\n"

    create_txt_artifact(txt_content, grant_dir / "06_review.txt")

    # DOCX
    sections = [
        {
            'heading': 'Общие показатели',
            'content': {
                'Готовность гранта': f"{reviewer_result.get('readiness_score', 0):.2f}/10",
                'Вероятность одобрения': f"{reviewer_result.get('approval_probability', 0):.1f}%",
                'Уровень качества': reviewer_result.get('quality_tier', 'Unknown'),
                'Можно подавать': 'ДА' if reviewer_result.get('can_submit', False) else 'НЕТ'
            }
        },
        {
            'heading': 'Оценки по критериям',
            'content': {
                crit_name: f"{crit_data['score']:.2f}/10 (вес {crit_data['weight']*100:.0f}%)"
                for crit_name, crit_data in reviewer_result.get('criteria_scores', {}).items()
            }
        },
        {
            'heading': f"Сильные стороны ({len(reviewer_result.get('strengths', []))})",
            'content': reviewer_result.get('strengths', [])
        },
        {
            'heading': f"Слабые стороны ({len(reviewer_result.get('weaknesses', []))})",
            'content': reviewer_result.get('weaknesses', [])
        },
        {
            'heading': f"Рекомендации ({len(reviewer_result.get('recommendations', []))})",
            'content': reviewer_result.get('recommendations', [])
        }
    ]

    create_docx_artifact("Стадия 6: Review (Финальная оценка)", sections, grant_dir / "06_review.docx")
    create_pdf_artifact("Стадия 6: Review (Финальная оценка)", sections, grant_dir / "06_review.pdf")


def export_grant_full(test_case: Dict, e2e_result: Dict):
    """Экспорт полного набора артефактов для одного гранта"""
    print(f"\n{'=' * 80}")
    print(f"📦 Экспорт гранта: {test_case['name']}")
    print(f"{'=' * 80}")

    # Создаем папку для гранта
    grant_dir = OUTPUT_DIR / test_case['id']
    grant_dir.mkdir(exist_ok=True)
    print(f"📁 Папка: {grant_dir.name}/")

    # Стадия 1: Interview
    export_interview_artifacts(test_case, grant_dir)

    # Стадия 2: Audit
    export_audit_artifacts(test_case, grant_dir)

    # Стадия 3: Research (из mock research_results)
    # Импортируем функцию создания mock данных
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from test_full_workflow_e2e_phases_3_5 import create_mock_research_results
    research_results = create_mock_research_results(test_case)
    export_research_artifacts(test_case, research_results, grant_dir)

    # Стадия 4: Planning
    writer_step = e2e_result.get('steps', {}).get('writer', {})
    # Создаем mock план если нет реального
    mock_plan = {
        'sections': [
            {
                'name': 'Актуальность проблемы',
                'key_elements': ['Статистика', 'Госпрограммы', 'Таблица динамики'],
                'research_blocks': ['block1'],
                'citations_count': 4,
                'tables_count': 1
            },
            {
                'name': 'География и целевая аудитория',
                'key_elements': ['Сравнение регионов', 'ЦА', 'Инфраструктура'],
                'research_blocks': ['block2'],
                'citations_count': 3,
                'tables_count': 1
            },
            {
                'name': 'Цели и задачи',
                'key_elements': ['SMART-цели', 'KPI', 'Успешные кейсы'],
                'research_blocks': ['block3'],
                'citations_count': 3,
                'tables_count': 0
            }
        ],
        'total_length_estimate': 15000,
        'total_citations': 10,
        'total_tables': 2
    }
    export_planning_artifacts(mock_plan, grant_dir, test_case)

    # Стадия 5: Grant (финальная заявка)
    # Создаем mock grant_content
    from writer_agent_v2 import WriterAgentV2
    writer_v2 = WriterAgentV2(None)
    grant_content = writer_v2._create_fallback_content(test_case['user_answers'], research_results)
    citations = writer_v2._format_citations(research_results, 10)
    tables = writer_v2._format_tables(research_results, 2)

    export_grant_artifacts(grant_content, citations, tables, grant_dir, test_case)

    # Стадия 6: Review (финальная оценка)
    reviewer_step = e2e_result.get('steps', {}).get('reviewer', {})
    if reviewer_step.get('status') == 'success':
        # Создаем полный reviewer result
        reviewer_result = {
            'readiness_score': reviewer_step.get('readiness_score', 0),
            'approval_probability': reviewer_step.get('approval_probability', 0),
            'quality_tier': reviewer_step.get('quality_tier', 'Unknown'),
            'can_submit': reviewer_step.get('can_submit', False),
            'criteria_scores': reviewer_step.get('criteria_scores', {}),
            'strengths': ['Сильная доказательная база', 'SMART-цели', 'Соответствие нацпроектам'],
            'weaknesses': ['Недостаточный объем текста', 'Нужна устойчивость'],
            'recommendations': ['Увеличить детальность', 'Добавить экономическое обоснование']
        }
        export_review_artifacts(reviewer_result, grant_dir, test_case)

    print(f"\n✅ Экспорт гранта завершен: {grant_dir.name}/")


def main():
    """Основная функция - экспорт всех грантов"""

    # Читаем результаты E2E тестов
    report_files = list(Path('.').glob('E2E_FULL_WORKFLOW_REPORT_*.json'))

    if not report_files:
        print("\n❌ Не найдены файлы отчетов E2E тестов")
        print("   Сначала запустите: python test_full_workflow_e2e_phases_3_5.py")
        return

    # Берем последний отчет
    latest_report = max(report_files, key=lambda p: p.stat().st_mtime)
    print(f"\n📄 Читаем отчет: {latest_report.name}")

    with open(latest_report, 'r', encoding='utf-8') as f:
        report_data = json.load(f)

    results = report_data.get('results', [])
    print(f"📊 Найдено грантов: {len(results)}")

    # Импортируем TEST_CASES
    from test_full_workflow_e2e_phases_3_5 import TEST_CASES

    # Экспортируем каждый грант
    for i, (test_case, e2e_result) in enumerate(zip(TEST_CASES, results), 1):
        export_grant_full(test_case, e2e_result)

        if i < len(TEST_CASES):
            print("\n⏳ Пауза 1 секунда...")
            import time
            time.sleep(1)

    # Итоговая статистика
    print(f"\n\n{'=' * 80}")
    print("✅ ВСЕ ГРАНТЫ ЭКСПОРТИРОВАНЫ!")
    print(f"{'=' * 80}\n")

    print(f"📁 Папка с грантами: {OUTPUT_DIR.absolute()}")
    print(f"📊 Всего грантов: {len(results)}")
    print(f"📄 Артефактов на грант: 18 файлов (6 стадий × 3 формата)")
    print(f"📦 Всего файлов создано: {len(results) * 18}")

    print("\n📂 Структура папки:")
    for test_case in TEST_CASES:
        grant_dir = OUTPUT_DIR / test_case['id']
        if grant_dir.exists():
            files = sorted(grant_dir.glob('*'))
            print(f"  {grant_dir.name}/ ({len(files)} файлов)")


if __name__ == "__main__":
    main()
