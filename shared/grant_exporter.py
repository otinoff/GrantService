#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Grant Exporter - Универсальный модуль экспорта грантовых заявок
Поддерживает форматы: MD, PDF, DOCX

Использование:
    from shared.grant_exporter import GrantExporter

    exporter = GrantExporter(grant_result_json)

    # Экспорт в файл
    exporter.export_to_markdown("output.md")
    exporter.export_to_pdf("output.pdf")
    exporter.export_to_docx("output.docx")

    # Экспорт в память (для отправки пользователю)
    md_content = exporter.get_markdown_content()
    pdf_bytes = exporter.get_pdf_bytes()
    docx_bytes = exporter.get_docx_bytes()

    # Экспорт всех форматов сразу
    files = exporter.export_all(output_dir="./output")
"""
import os
import json
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)


class GrantExporter:
    """Универсальный экспортер грантовых заявок в MD, PDF, DOCX"""

    def __init__(self, grant_data: Dict[str, Any]):
        """
        Инициализация экспортера

        Args:
            grant_data: JSON с результатом работы Writer Agent V2
                Должен содержать ключи: status, application, citations, tables, quality_score
        """
        self.grant_data = grant_data
        self.application = grant_data.get('application', {})
        self.citations = grant_data.get('citations', [])
        self.tables = grant_data.get('tables', [])
        self.quality_score = grant_data.get('quality_score', 0)
        self.provider_used = grant_data.get('provider_used', 'unknown')
        self.timestamp = grant_data.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # Метаданные
        self.metadata = self.application.get('metadata', {})
        self.total_chars = self.metadata.get('total_chars', 0)
        self.citations_used = self.metadata.get('citations_used', len(self.citations))
        self.tables_included = self.metadata.get('tables_included', len(self.tables))

        logger.info(f"✅ GrantExporter инициализирован: {self.total_chars} символов, "
                   f"{self.citations_used} цитат, {self.tables_included} таблиц")

    # ==================== MARKDOWN EXPORT ====================

    def get_markdown_content(self) -> str:
        """
        Создать содержимое Markdown файла

        Returns:
            str: Полный текст грантовой заявки в формате Markdown
        """
        md_lines = []

        # Заголовок
        title = self.application.get('title', 'Грантовая заявка')
        md_lines.append(f"# 📋 {title}\n")
        md_lines.append(f"**Дата создания:** {self.timestamp}\n")
        md_lines.append(f"**Оценка качества:** {self.quality_score}/10 ⭐\n")
        md_lines.append(f"**Объем:** {self.total_chars} символов\n")
        md_lines.append(f"**Цитаты:** {self.citations_used} | **Таблицы:** {self.tables_included}\n")
        md_lines.append(f"**LLM Provider:** {self.provider_used}\n")
        md_lines.append("\n---\n\n")

        # 9 разделов заявки
        sections = [
            ("1. КРАТКОЕ ОПИСАНИЕ ПРОЕКТА", "section_1_brief"),
            ("2. ОБОСНОВАНИЕ СОЦИАЛЬНОЙ ЗНАЧИМОСТИ", "section_2_problem"),
            ("3. ЦЕЛЬ ПРОЕКТА", "section_3_goal"),
            ("4. ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ", "section_4_results"),
            ("5. ЗАДАЧИ ПРОЕКТА", "section_5_tasks"),
            ("6. ПАРТНЕРЫ ПРОЕКТА", "section_6_partners"),
            ("7. ИНФОРМАЦИОННОЕ СОПРОВОЖДЕНИЕ", "section_7_info"),
            ("8. ДАЛЬНЕЙШЕЕ РАЗВИТИЕ ПРОЕКТА", "section_8_future"),
            ("9. КАЛЕНДАРНЫЙ ПЛАН", "section_9_calendar")
        ]

        for section_title, section_key in sections:
            md_lines.append(f"## {section_title}\n\n")
            content = self.application.get(section_key, '')
            if content:
                md_lines.append(f"{content}\n\n")
            else:
                md_lines.append("*Нет данных*\n\n")
            md_lines.append("---\n\n")

        # Приложение 1: Цитаты
        md_lines.append("## 📚 ПРИЛОЖЕНИЕ 1: ЦИТАТЫ И ИСТОЧНИКИ\n\n")
        if self.citations:
            for i, citation in enumerate(self.citations, 1):
                md_lines.append(f"### Цитата {i}\n\n")
                md_lines.append(f"**Тип:** {citation.get('type', 'Не указан')}\n\n")
                md_lines.append(f"> {citation.get('text', '')}\n\n")
                if citation.get('source'):
                    md_lines.append(f"**Источник:** {citation['source']}")
                    if citation.get('date'):
                        md_lines.append(f" ({citation['date']})")
                    md_lines.append("\n\n")
        else:
            md_lines.append("*Цитаты отсутствуют*\n\n")

        md_lines.append("---\n\n")

        # Приложение 2: Таблицы
        md_lines.append("## 📊 ПРИЛОЖЕНИЕ 2: ТАБЛИЦЫ И ДАННЫЕ\n\n")
        if self.tables:
            for i, table in enumerate(self.tables, 1):
                md_lines.append(f"### Таблица {i}: {table.get('title', 'Без названия')}\n\n")
                md_lines.append(f"**Тип:** {table.get('type', 'Не указан')}\n\n")
                md_lines.append(f"**Источник:** {table.get('source', 'Не указан')}\n\n")

                # Если есть данные таблицы - показать первые 5 полей
                table_data = table.get('data', {})
                if isinstance(table_data, dict) and table_data:
                    md_lines.append("**Данные:**\n\n")
                    for key, value in list(table_data.items())[:5]:
                        md_lines.append(f"- **{key}:** {str(value)[:100]}\n")
                    md_lines.append("\n")
        else:
            md_lines.append("*Таблицы отсутствуют*\n\n")

        # Футер
        md_lines.append("---\n\n")
        md_lines.append(f"*Сгенерировано Writer Agent V2 | {self.timestamp}*\n")

        return ''.join(md_lines)

    def export_to_markdown(self, output_path: str) -> str:
        """
        Экспортировать в Markdown файл

        Args:
            output_path: Путь к выходному файлу

        Returns:
            str: Путь к созданному файлу
        """
        try:
            md_content = self.get_markdown_content()

            # Создаем директорию если не существует
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)

            logger.info(f"✅ MD файл создан: {output_path} ({len(md_content)} символов)")
            return output_path
        except Exception as e:
            logger.error(f"❌ Ошибка создания MD: {e}")
            raise

    # ==================== PDF EXPORT ====================

    def get_pdf_bytes(self) -> bytes:
        """
        Создать PDF в памяти (для отправки пользователю)

        Returns:
            bytes: PDF файл в виде байтов
        """
        # Стратегия 1: Попробовать Pandoc (для Windows тестирования)
        try:
            pdf_bytes = self._create_pdf_with_pandoc()
            logger.info(f"✅ PDF создан через Pandoc: {len(pdf_bytes)} байт")
            return pdf_bytes
        except Exception as e:
            logger.warning(f"⚠️ Pandoc недоступен: {e}")

        # Стратегия 2: Попробовать WeasyPrint (для Linux продакшн)
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration

            # Получаем MD контент
            md_content = self.get_markdown_content()

            # Конвертируем MD в HTML (простая конвертация)
            html_content = self._markdown_to_html(md_content)

            # CSS для красивого форматирования
            css = CSS(string='''
                @page { size: A4; margin: 2cm; }
                body { font-family: "DejaVu Sans", Arial, sans-serif; font-size: 11pt; line-height: 1.6; }
                h1 { color: #2c3e50; font-size: 24pt; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
                h2 { color: #34495e; font-size: 18pt; margin-top: 20px; border-bottom: 2px solid #95a5a6; padding-bottom: 5px; }
                h3 { color: #7f8c8d; font-size: 14pt; margin-top: 15px; }
                table { width: 100%; border-collapse: collapse; margin: 10px 0; }
                th, td { border: 1px solid #bdc3c7; padding: 8px; text-align: left; }
                th { background-color: #ecf0f1; font-weight: bold; }
                blockquote { border-left: 4px solid #3498db; padding-left: 15px; margin-left: 0; color: #555; }
                strong { color: #2c3e50; }
                hr { border: none; border-top: 1px solid #bdc3c7; margin: 20px 0; }
            ''')

            font_config = FontConfiguration()

            # Генерируем PDF в память
            pdf_bytes = HTML(string=html_content).write_pdf(
                stylesheets=[css],
                font_config=font_config
            )

            logger.info(f"✅ PDF создан через WeasyPrint: {len(pdf_bytes)} байт")
            return pdf_bytes
        except ImportError:
            logger.warning("⚠️ WeasyPrint не установлен")
        except Exception as e:
            logger.warning(f"⚠️ WeasyPrint ошибка: {e}")

        # Стратегия 3: Fallback на ReportLab
        logger.warning("⚠️ Используем упрощенную версию PDF через ReportLab")
        return self._create_simple_pdf()

    def export_to_pdf(self, output_path: str) -> str:
        """
        Экспортировать в PDF файл

        Args:
            output_path: Путь к выходному файлу

        Returns:
            str: Путь к созданному файлу
        """
        try:
            pdf_bytes = self.get_pdf_bytes()

            # Создаем директорию если не существует
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)

            logger.info(f"✅ PDF файл создан: {output_path} ({len(pdf_bytes)} байт)")
            return output_path
        except Exception as e:
            logger.error(f"❌ Ошибка экспорта PDF: {e}")
            raise

    def _create_pdf_with_pandoc(self) -> bytes:
        """
        Создать PDF через Pandoc (для Windows тестирования)

        Returns:
            bytes: PDF файл

        Raises:
            Exception: Если Pandoc недоступен
        """
        import subprocess
        import tempfile

        # Проверяем наличие Pandoc
        pandoc_paths = [
            r"C:\Program Files\Pandoc\pandoc.exe",
            r"C:\Program Files (x86)\Pandoc\pandoc.exe",
            "pandoc"  # в PATH
        ]

        pandoc_exe = None
        for path in pandoc_paths:
            try:
                result = subprocess.run([path, "--version"], capture_output=True, timeout=5)
                if result.returncode == 0:
                    pandoc_exe = path
                    logger.info(f"✅ Pandoc найден: {path}")
                    break
            except:
                continue

        if not pandoc_exe:
            raise FileNotFoundError("Pandoc не найден. Установите: https://pandoc.org/")

        # Создаем временный MD файл
        md_content = self.get_markdown_content()

        with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', suffix='.md', delete=False) as md_file:
            md_file.write(md_content)
            md_path = md_file.name

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
            pdf_path = pdf_file.name

        try:
            # Стратегия: MD -> HTML -> самодельный PDF (без LaTeX)
            # Создаем HTML через Pandoc с красивым оформлением
            html_path = pdf_path.replace('.pdf', '.html')

            cmd_html = [
                pandoc_exe,
                md_path,
                "-o", html_path,
                "--standalone",
                "--self-contained",
                "-c", "https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css",
                "--metadata", "title=Грантовая заявка"
            ]

            result = subprocess.run(cmd_html, capture_output=True, timeout=30, text=True)

            if result.returncode != 0:
                raise Exception(f"Pandoc HTML ошибка: {result.stderr}")

            # Теперь используем wkhtmltopdf если есть, иначе fallback
            wkhtmltopdf_paths = [
                r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe",
                r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe",
                "wkhtmltopdf"
            ]

            wkhtmltopdf_exe = None
            for wk_path in wkhtmltopdf_paths:
                try:
                    test = subprocess.run([wk_path, "--version"], capture_output=True, timeout=5)
                    if test.returncode == 0:
                        wkhtmltopdf_exe = wk_path
                        break
                except:
                    continue

            if wkhtmltopdf_exe:
                # Конвертируем HTML -> PDF через wkhtmltopdf
                logger.info("✅ Используем wkhtmltopdf для PDF")
                cmd_pdf = [wkhtmltopdf_exe, html_path, pdf_path, "--enable-local-file-access"]
                result = subprocess.run(cmd_pdf, capture_output=True, timeout=30)

                if result.returncode != 0:
                    raise Exception("wkhtmltopdf ошибка")
            else:
                # Fallback: читаем HTML и используем weasyprint/reportlab
                logger.warning("⚠️ wkhtmltopdf не найден, используем fallback")
                with open(html_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()

                # Пробуем weasyprint
                try:
                    from weasyprint import HTML
                    HTML(string=html_content).write_pdf(pdf_path)
                except:
                    # Последний fallback - создаем простой PDF
                    raise Exception("Невозможно создать PDF - нет подходящих инструментов")

            # Читаем PDF
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()

            return pdf_bytes

        finally:
            # Удаляем временные файлы
            try:
                os.unlink(md_path)
                os.unlink(pdf_path)
            except:
                pass

    def _markdown_to_html(self, md_content: str) -> str:
        """Простая конвертация Markdown в HTML"""
        try:
            import markdown
            html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        except ImportError:
            # Fallback: простая замена без markdown библиотеки
            html_body = md_content.replace('\n\n', '</p><p>').replace('\n', '<br>')
            html_body = f"<p>{html_body}</p>"

        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Грантовая заявка</title>
        </head>
        <body>
            {html_body}
        </body>
        </html>
        """

    def _create_simple_pdf(self) -> bytes:
        """Упрощенная версия PDF через reportlab (fallback)"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import cm

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)

            # Регистрируем шрифт с поддержкой кириллицы
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
                font_name = 'DejaVuSans'
            except:
                font_name = 'Helvetica'

            styles = getSampleStyleSheet()

            # Кастомные стили
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                         fontSize=18, textColor=colors.HexColor('#2c3e50'),
                                         spaceAfter=12, fontName=font_name)
            heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                          fontSize=14, textColor=colors.HexColor('#34495e'),
                                          spaceAfter=10, spaceBefore=10, fontName=font_name)
            normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'],
                                         fontSize=11, fontName=font_name)

            story = []

            # Заголовок
            title = self.application.get('title', 'Грантовая заявка')
            story.append(Paragraph(f"📋 {title}", title_style))
            story.append(Spacer(1, 0.5*cm))

            # Метаданные
            story.append(Paragraph(f"<b>Дата:</b> {self.timestamp}", normal_style))
            story.append(Paragraph(f"<b>Оценка:</b> {self.quality_score}/10", normal_style))
            story.append(Paragraph(f"<b>Объем:</b> {self.total_chars} символов", normal_style))
            story.append(Spacer(1, 0.5*cm))

            # Разделы (упрощенно - только первые 3)
            sections = [
                ("1. КРАТКОЕ ОПИСАНИЕ", "section_1_brief"),
                ("2. ОБОСНОВАНИЕ ПРОБЛЕМЫ", "section_2_problem"),
                ("3. ЦЕЛЬ ПРОЕКТА", "section_3_goal")
            ]

            for section_title, section_key in sections:
                story.append(Paragraph(section_title, heading_style))
                content = self.application.get(section_key, '')[:2000]  # Ограничим для простоты
                if content:
                    # Разбиваем на параграфы
                    paragraphs = content.split('\n\n')
                    for para in paragraphs[:5]:  # Макс 5 параграфов
                        if para.strip():
                            story.append(Paragraph(para.strip(), normal_style))
                            story.append(Spacer(1, 0.3*cm))
                story.append(PageBreak())

            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()

            logger.info(f"✅ Простой PDF создан: {len(pdf_bytes)} байт")
            return pdf_bytes
        except Exception as e:
            logger.error(f"❌ Ошибка создания простого PDF: {e}")
            raise

    # ==================== DOCX EXPORT ====================

    def get_docx_bytes(self) -> bytes:
        """
        Создать DOCX в памяти (для отправки пользователю)

        Returns:
            bytes: DOCX файл в виде байтов
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document()

            # Заголовок
            title = self.application.get('title', 'Грантовая заявка')
            heading = doc.add_heading(f"📋 {title}", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Метаданные
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Дата создания: ").bold = True
            meta_para.add_run(f"{self.timestamp}\n")
            meta_para.add_run(f"Оценка качества: ").bold = True
            meta_para.add_run(f"{self.quality_score}/10 ⭐\n")
            meta_para.add_run(f"Объем: ").bold = True
            meta_para.add_run(f"{self.total_chars} символов\n")
            meta_para.add_run(f"Цитаты: ").bold = True
            meta_para.add_run(f"{self.citations_used} | ")
            meta_para.add_run(f"Таблицы: ").bold = True
            meta_para.add_run(f"{self.tables_included}\n")

            doc.add_paragraph()  # Пустая строка
            doc.add_paragraph("_" * 80)  # Разделитель
            doc.add_paragraph()

            # 9 разделов
            sections = [
                ("1. КРАТКОЕ ОПИСАНИЕ ПРОЕКТА", "section_1_brief"),
                ("2. ОБОСНОВАНИЕ СОЦИАЛЬНОЙ ЗНАЧИМОСТИ", "section_2_problem"),
                ("3. ЦЕЛЬ ПРОЕКТА", "section_3_goal"),
                ("4. ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ", "section_4_results"),
                ("5. ЗАДАЧИ ПРОЕКТА", "section_5_tasks"),
                ("6. ПАРТНЕРЫ ПРОЕКТА", "section_6_partners"),
                ("7. ИНФОРМАЦИОННОЕ СОПРОВОЖДЕНИЕ", "section_7_info"),
                ("8. ДАЛЬНЕЙШЕЕ РАЗВИТИЕ ПРОЕКТА", "section_8_future"),
                ("9. КАЛЕНДАРНЫЙ ПЛАН", "section_9_calendar")
            ]

            for section_title, section_key in sections:
                # Заголовок раздела
                section_heading = doc.add_heading(section_title, level=2)
                run = section_heading.runs[0]
                run.font.color.rgb = RGBColor(52, 73, 94)  # Темно-синий

                # Содержимое раздела
                content = self.application.get(section_key, '')
                if content:
                    # Разбиваем на параграфы
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.paragraph_format.line_spacing = 1.5
                else:
                    doc.add_paragraph("Нет данных").italic = True

                doc.add_paragraph()  # Пустая строка
                doc.add_paragraph("_" * 80)  # Разделитель
                doc.add_paragraph()

            # Приложение: Цитаты
            if self.citations:
                doc.add_heading("📚 ПРИЛОЖЕНИЕ 1: ЦИТАТЫ И ИСТОЧНИКИ", level=2)
                for i, citation in enumerate(self.citations, 1):
                    doc.add_heading(f"Цитата {i}", level=3)
                    cite_para = doc.add_paragraph()
                    cite_para.add_run(f"Тип: ").bold = True
                    cite_para.add_run(f"{citation.get('type', 'Не указан')}\n")
                    cite_para.add_run(f"Текст: ").bold = True
                    cite_para.add_run(f'"{citation.get("text", "")}"').italic = True
                    cite_para.add_run(f"\nИсточник: ").bold = True
                    source_text = citation.get('source', 'Не указан')
                    if citation.get('date'):
                        source_text += f" ({citation['date']})"
                    cite_para.add_run(source_text)
                    doc.add_paragraph()

            # Сохраняем в BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            logger.info(f"✅ DOCX создан в памяти: {len(docx_bytes)} байт")
            return docx_bytes
        except ImportError:
            logger.error("❌ python-docx не установлен. Установите: pip install python-docx")
            raise ImportError("python-docx library is required for DOCX export")
        except Exception as e:
            logger.error(f"❌ Ошибка создания DOCX: {e}")
            raise

    def export_to_docx(self, output_path: str) -> str:
        """
        Экспортировать в DOCX файл

        Args:
            output_path: Путь к выходному файлу

        Returns:
            str: Путь к созданному файлу
        """
        try:
            docx_bytes = self.get_docx_bytes()

            # Создаем директорию если не существует
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            with open(output_path, 'wb') as f:
                f.write(docx_bytes)

            logger.info(f"✅ DOCX файл создан: {output_path} ({len(docx_bytes)} байт)")
            return output_path
        except Exception as e:
            logger.error(f"❌ Ошибка экспорта DOCX: {e}")
            raise

    # ==================== UNIVERSAL EXPORT ====================

    def export_all(self, output_dir: str, base_filename: Optional[str] = None) -> Dict[str, str]:
        """
        Экспортировать во все форматы сразу

        Args:
            output_dir: Директория для сохранения файлов
            base_filename: Базовое имя файла (без расширения).
                          Если не указано, используется application_number или timestamp

        Returns:
            Dict[str, str]: Словарь с путями к созданным файлам
                {
                    'markdown': '/path/to/file.md',
                    'pdf': '/path/to/file.pdf',
                    'docx': '/path/to/file.docx'
                }
        """
        try:
            # Определяем базовое имя файла
            if not base_filename:
                app_number = self.grant_data.get('application_number',
                                                 f"grant_{self.timestamp.replace(':', '-').replace(' ', '_')}")
                base_filename = app_number.replace('#', '').replace('/', '_')

            # Создаем директорию
            os.makedirs(output_dir, exist_ok=True)

            results = {}

            # Экспорт в MD
            try:
                md_path = os.path.join(output_dir, f"{base_filename}.md")
                results['markdown'] = self.export_to_markdown(md_path)
            except Exception as e:
                logger.error(f"❌ Ошибка экспорта MD: {e}")
                results['markdown'] = None

            # Экспорт в PDF
            try:
                pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")
                results['pdf'] = self.export_to_pdf(pdf_path)
            except Exception as e:
                logger.error(f"❌ Ошибка экспорта PDF: {e}")
                results['pdf'] = None

            # Экспорт в DOCX
            try:
                docx_path = os.path.join(output_dir, f"{base_filename}.docx")
                results['docx'] = self.export_to_docx(docx_path)
            except Exception as e:
                logger.error(f"❌ Ошибка экспорта DOCX: {e}")
                results['docx'] = None

            # Логируем результаты
            success_count = sum(1 for v in results.values() if v is not None)
            logger.info(f"✅ Экспорт завершен: {success_count}/3 форматов успешно")

            return results
        except Exception as e:
            logger.error(f"❌ Критическая ошибка export_all: {e}")
            raise


# ==================== UTILITY FUNCTIONS ====================

def export_grant_from_json_file(json_path: str, output_dir: str,
                                base_filename: Optional[str] = None) -> Dict[str, str]:
    """
    Удобная функция для экспорта из JSON файла

    Args:
        json_path: Путь к JSON файлу с результатом Writer Agent
        output_dir: Директория для сохранения файлов
        base_filename: Базовое имя файла (опционально)

    Returns:
        Dict[str, str]: Словарь с путями к созданным файлам

    Example:
        >>> from shared.grant_exporter import export_grant_from_json_file
        >>> files = export_grant_from_json_file(
        ...     "grant_export_session_9/03_grant_draft/grant_result.json",
        ...     "grant_export_session_9/03_grant_draft/"
        ... )
        >>> print(files['markdown'])
        '/path/to/grant.md'
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            grant_data = json.load(f)

        exporter = GrantExporter(grant_data)
        return exporter.export_all(output_dir, base_filename)
    except Exception as e:
        logger.error(f"❌ Ошибка экспорта из JSON файла {json_path}: {e}")
        raise


if __name__ == "__main__":
    # Пример использования
    print("GrantExporter - Универсальный модуль экспорта грантовых заявок")
    print("Используйте:")
    print("  from shared.grant_exporter import GrantExporter, export_grant_from_json_file")
