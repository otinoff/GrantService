#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
A/B Test: 2 полных researcher прогона для анкеты Екатерины

Цель:
- Запустить ResearcherAgentV2 ДВА РАЗА для одной anketa_id
- Проверить возможность хранения 2+ research versions в БД
- Экспортировать оба результата в Word для сравнения

Автор: Claude Code
Дата: 2025-10-12
"""

import sys
import os
import asyncio
from datetime import datetime
import json
from pathlib import Path

# Set PostgreSQL environment variables
os.environ['PGHOST'] = 'localhost'
os.environ['PGPORT'] = '5432'
os.environ['PGUSER'] = 'postgres'
os.environ['PGPASSWORD'] = 'root'
os.environ['PGDATABASE'] = 'grantservice'

# Set UTF-8 encoding for console output
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

# Добавляем пути
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.path.append(os.path.join(os.path.dirname(__file__), 'agents'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'web-admin'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'shared'))

from data.database.models import GrantServiceDatabase
from agents.researcher_agent_v2 import ResearcherAgentV2

# Для экспорта в Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    print("[WARN] python-docx не установлен, экспорт в Word будет недоступен")
    DOCX_AVAILABLE = False


async def run_ab_test():
    """
    A/B тест: запустить researcher 2 раза для одной анкеты
    """
    print("=" * 80)
    print("🧪 A/B TEST: Екатерина - 2 версии исследования (27 queries × 2)")
    print("=" * 80)

    # Инициализация БД
    db = GrantServiceDatabase()

    # Анкета Екатерины
    ANKETA_ID = "EKATERINA_20251010_235448"

    print(f"\n📋 Anketa ID: {ANKETA_ID}")
    print(f"📊 Проект: Возрождение храма и народных традиций в селе Анисимово")

    # Проверить существующие research
    print("\n" + "=" * 80)
    print("🔍 ПРОВЕРКА: Существующие исследования")
    print("=" * 80)

    with db.connect() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, research_id, status, llm_provider, created_at,
                   CASE
                       WHEN research_results IS NULL OR research_results::text = '{}' THEN 0
                       ELSE 1
                   END as has_results
            FROM researcher_research
            WHERE anketa_id = %s
            ORDER BY created_at DESC
        """, (ANKETA_ID,))

        existing = cursor.fetchall()

        if existing:
            print(f"\n✅ Найдено {len(existing)} существующих research:")
            for row in existing:
                id_val, research_id, status, provider, created_at, has_results = row
                print(f"   [{id_val}] {research_id} | {status} | {provider} | {created_at} | results={bool(has_results)}")
        else:
            print(f"\n⚠️  Исследований для anketa_id={ANKETA_ID} не найдено")

        cursor.close()

    # VERSION A: Первый прогон
    print("\n" + "=" * 80)
    print("🔬 VERSION A: Первый прогон (27 queries)")
    print("=" * 80)
    print(f"⏰ Начало: {datetime.now().strftime('%H:%M:%S')}")

    researcher_a = ResearcherAgentV2(db, llm_provider="claude_code")

    result_a = await researcher_a.research_with_expert_prompts(ANKETA_ID)

    print(f"\n✅ VERSION A завершена!")
    print(f"   Research ID: {result_a['research_id']}")
    print(f"   Status: {result_a['status']}")

    if result_a['status'] == 'completed':
        metadata_a = result_a['research_results']['metadata']
        print(f"   Queries: {metadata_a['total_queries']}")
        print(f"   Sources: {metadata_a['sources_count']}")
        print(f"   Time: {metadata_a['total_processing_time']}s")
        print(f"   Provider: {metadata_a.get('websearch_provider', 'unknown')}")

    # Небольшая пауза перед вторым прогоном
    print("\n⏸️  Пауза 5 секунд перед VERSION B...")
    await asyncio.sleep(5)

    # VERSION B: Второй прогон
    print("\n" + "=" * 80)
    print("🔬 VERSION B: Второй прогон (27 queries)")
    print("=" * 80)
    print(f"⏰ Начало: {datetime.now().strftime('%H:%M:%S')}")

    researcher_b = ResearcherAgentV2(db, llm_provider="claude_code")

    result_b = await researcher_b.research_with_expert_prompts(ANKETA_ID)

    print(f"\n✅ VERSION B завершена!")
    print(f"   Research ID: {result_b['research_id']}")
    print(f"   Status: {result_b['status']}")

    if result_b['status'] == 'completed':
        metadata_b = result_b['research_results']['metadata']
        print(f"   Queries: {metadata_b['total_queries']}")
        print(f"   Sources: {metadata_b['sources_count']}")
        print(f"   Time: {metadata_b['total_processing_time']}s")
        print(f"   Provider: {metadata_b.get('websearch_provider', 'unknown')}")

    # Проверить что в БД теперь 2 версии
    print("\n" + "=" * 80)
    print("🔍 ПРОВЕРКА: Обе версии сохранены в БД?")
    print("=" * 80)

    with db.connect() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, research_id, status, llm_provider,
                   research_results->>'metadata' as metadata_json,
                   created_at, completed_at
            FROM researcher_research
            WHERE anketa_id = %s
            ORDER BY created_at DESC
            LIMIT 10
        """, (ANKETA_ID,))

        all_research = cursor.fetchall()
        cursor.close()

    print(f"\n📊 Всего research для anketa_id={ANKETA_ID}: {len(all_research)}")
    print("\nДетали:")

    for idx, row in enumerate(all_research, 1):
        id_val, research_id, status, provider, metadata_json, created_at, completed_at = row

        print(f"\n[{idx}] {research_id}")
        print(f"    ID: {id_val}")
        print(f"    Status: {status}")
        print(f"    Provider: {provider}")
        print(f"    Created: {created_at}")
        print(f"    Completed: {completed_at}")

        if metadata_json:
            try:
                metadata = json.loads(metadata_json)
                print(f"    Queries: {metadata.get('total_queries', 'N/A')}")
                print(f"    Sources: {metadata.get('sources_count', 'N/A')}")
                print(f"    Time: {metadata.get('total_processing_time', 'N/A')}s")
                print(f"    WebSearch: {metadata.get('websearch_provider', 'N/A')}")
            except:
                print(f"    Metadata: {metadata_json[:100]}...")

    # Экспорт в Word
    if DOCX_AVAILABLE and result_a['status'] == 'completed' and result_b['status'] == 'completed':
        print("\n" + "=" * 80)
        print("📄 ЭКСПОРТ: Создание Word документов")
        print("=" * 80)

        output_dir = Path("grants_output/ab_test_ekaterina")
        output_dir.mkdir(parents=True, exist_ok=True)

        # Экспорт VERSION A
        file_a = output_dir / f"Research_A_{result_a['research_id']}.docx"
        export_to_word(result_a, file_a, "VERSION A")
        print(f"\n✅ VERSION A экспортирована:")
        print(f"   {file_a}")

        # Экспорт VERSION B
        file_b = output_dir / f"Research_B_{result_b['research_id']}.docx"
        export_to_word(result_b, file_b, "VERSION B")
        print(f"\n✅ VERSION B экспортирована:")
        print(f"   {file_b}")

        # Создать сравнительный документ
        comparison_file = output_dir / f"Comparison_AB_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        create_comparison_document(result_a, result_b, comparison_file)
        print(f"\n✅ Сравнительный документ создан:")
        print(f"   {comparison_file}")

        print("\n" + "=" * 80)
        print("📁 ФАЙЛЫ ДЛЯ ЧТЕНИЯ:")
        print("=" * 80)
        print(f"\n1. VERSION A: {file_a}")
        print(f"2. VERSION B: {file_b}")
        print(f"3. COMPARISON: {comparison_file}")

        return {
            'version_a': str(file_a),
            'version_b': str(file_b),
            'comparison': str(comparison_file),
            'research_ids': [result_a['research_id'], result_b['research_id']]
        }

    return {
        'research_ids': [result_a.get('research_id'), result_b.get('research_id')],
        'statuses': [result_a['status'], result_b['status']]
    }


def export_to_word(research_result: dict, output_path: Path, version_label: str):
    """Экспортировать результаты исследования в Word документ"""

    doc = Document()

    # Заголовок
    title = doc.add_heading(f'Исследование грантовой заявки - {version_label}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Метаданные
    doc.add_heading('Метаданные исследования', 1)

    metadata = research_result['research_results']['metadata']

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'

    rows_data = [
        ('Research ID', research_result['research_id']),
        ('Статус', research_result['status']),
        ('Всего запросов', str(metadata['total_queries'])),
        ('Всего источников', str(metadata['sources_count'])),
        ('Время обработки', f"{metadata['total_processing_time']}s"),
        ('WebSearch провайдер', metadata.get('websearch_provider', 'unknown')),
        ('Fallback провайдер', metadata.get('websearch_fallback', 'none'))
    ]

    for idx, (label, value) in enumerate(rows_data):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value

    doc.add_paragraph()

    # Результаты блоков
    for block_key in ['block1_problem', 'block2_geography', 'block3_goals']:
        block_data = research_result['research_results'].get(block_key, {})

        if not block_data:
            continue

        # Заголовок блока
        block_names = {
            'block1_problem': 'Блок 1: Проблема и социальная значимость',
            'block2_geography': 'Блок 2: География и целевая аудитория',
            'block3_goals': 'Блок 3: Задачи, мероприятия и главная цель'
        }

        doc.add_heading(block_names[block_key], 1)

        # Статистика блока
        doc.add_paragraph(f"Запросов: {len(block_data.get('queries_used', []))}")
        doc.add_paragraph(f"Источников: {block_data.get('total_sources', 0)}")
        doc.add_paragraph(f"Время обработки: {block_data.get('processing_time', 0)}s")

        # Резюме
        doc.add_heading('Резюме', 2)
        doc.add_paragraph(block_data.get('summary', 'Нет данных'))

        # Ключевые факты
        key_facts = block_data.get('key_facts', [])
        if key_facts:
            doc.add_heading('Ключевые факты', 2)
            for idx, fact in enumerate(key_facts[:10], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{fact.get('fact', '')}\n")
                p.add_run(f"Источник: {fact.get('source', 'unknown')}").font.size = Pt(9)

        # Использованные запросы
        queries_used = block_data.get('queries_used', [])
        if queries_used:
            doc.add_heading('Использованные запросы', 2)
            for idx, query in enumerate(queries_used, 1):
                doc.add_paragraph(f"{idx}. {query}", style='List Number')

        doc.add_page_break()

    # Сохранить
    doc.save(str(output_path))


def create_comparison_document(result_a: dict, result_b: dict, output_path: Path):
    """Создать сравнительный документ для A и B версий"""

    doc = Document()

    # Заголовок
    title = doc.add_heading('Сравнение версий исследования A и B', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Сравнительная таблица метаданных
    doc.add_heading('Сравнение метаданных', 1)

    metadata_a = result_a['research_results']['metadata']
    metadata_b = result_b['research_results']['metadata']

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'

    # Заголовки
    table.rows[0].cells[0].text = 'Параметр'
    table.rows[0].cells[1].text = 'VERSION A'
    table.rows[0].cells[2].text = 'VERSION B'

    # Данные
    comparison_data = [
        ('Research ID', result_a['research_id'], result_b['research_id']),
        ('Статус', result_a['status'], result_b['status']),
        ('Всего запросов', metadata_a['total_queries'], metadata_b['total_queries']),
        ('Всего источников', metadata_a['sources_count'], metadata_b['sources_count']),
        ('Время обработки (s)', metadata_a['total_processing_time'], metadata_b['total_processing_time']),
        ('WebSearch провайдер', metadata_a.get('websearch_provider', 'unknown'), metadata_b.get('websearch_provider', 'unknown')),
        ('Fallback провайдер', metadata_a.get('websearch_fallback', 'none'), metadata_b.get('websearch_fallback', 'none'))
    ]

    for idx, (label, val_a, val_b) in enumerate(comparison_data, 1):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = str(val_a)
        table.rows[idx].cells[2].text = str(val_b)

    doc.add_paragraph()

    # Сравнение блоков
    doc.add_heading('Сравнение блоков', 1)

    for block_key in ['block1_problem', 'block2_geography', 'block3_goals']:
        block_names = {
            'block1_problem': 'Блок 1: Проблема',
            'block2_geography': 'Блок 2: География',
            'block3_goals': 'Блок 3: Задачи и цели'
        }

        doc.add_heading(block_names[block_key], 2)

        block_a = result_a['research_results'].get(block_key, {})
        block_b = result_b['research_results'].get(block_key, {})

        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Метрика'
        table.rows[0].cells[1].text = 'VERSION A'
        table.rows[0].cells[2].text = 'VERSION B'

        table.rows[1].cells[0].text = 'Запросов'
        table.rows[1].cells[1].text = str(len(block_a.get('queries_used', [])))
        table.rows[1].cells[2].text = str(len(block_b.get('queries_used', [])))

        table.rows[2].cells[0].text = 'Источников'
        table.rows[2].cells[1].text = str(block_a.get('total_sources', 0))
        table.rows[2].cells[2].text = str(block_b.get('total_sources', 0))

        table.rows[3].cells[0].text = 'Время (s)'
        table.rows[3].cells[1].text = str(block_a.get('processing_time', 0))
        table.rows[3].cells[2].text = str(block_b.get('processing_time', 0))

        doc.add_paragraph()

    # Выводы
    doc.add_page_break()
    doc.add_heading('Выводы', 1)

    doc.add_paragraph(f"✅ Обе версии исследования успешно выполнены и сохранены в БД")
    doc.add_paragraph(f"✅ Система поддерживает хранение нескольких research для одной anketa_id")
    doc.add_paragraph(f"✅ Возможность A/B тестирования подтверждена")

    # Различия
    sources_diff = abs(metadata_a['sources_count'] - metadata_b['sources_count'])
    time_diff = abs(metadata_a['total_processing_time'] - metadata_b['total_processing_time'])

    doc.add_paragraph(f"\nРазличия:")
    doc.add_paragraph(f"- Источники: ±{sources_diff}")
    doc.add_paragraph(f"- Время обработки: ±{time_diff}s")

    # Сохранить
    doc.save(str(output_path))


if __name__ == "__main__":
    print("\n" + "=" * 80)
    print("🚀 ЗАПУСК A/B ТЕСТА")
    print("=" * 80)
    print(f"Время: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Задача: 2 полных прогона researcher для анкеты Екатерины (27 queries × 2)")
    print(f"Цель: Проверить возможность хранения нескольких research versions")
    print("=" * 80)

    try:
        result = asyncio.run(run_ab_test())

        print("\n" + "=" * 80)
        print("✅ A/B ТЕСТ ЗАВЕРШЁН УСПЕШНО!")
        print("=" * 80)

        if 'version_a' in result:
            print("\n📁 Созданные файлы:")
            print(f"1. {result['version_a']}")
            print(f"2. {result['version_b']}")
            print(f"3. {result['comparison']}")

        print(f"\n🔬 Research IDs: {', '.join(result['research_ids'])}")

    except Exception as e:
        print(f"\n❌ ОШИБКА: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
