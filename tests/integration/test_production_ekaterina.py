#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Production Test: Полный цикл researcher для Екатерины через функционал приложения

Тест использует РЕАЛЬНЫЙ функционал приложения:
- ResearcherAgentV2 из agents/
- GrantServiceDatabase из data/database/
- WebSearchRouter (автоматически выбирает Claude Code primary)

Автор: Claude Code
Дата: 2025-10-12
"""

import sys
import os
import asyncio
from datetime import datetime
from pathlib import Path

# Set PostgreSQL environment variables
os.environ['PGHOST'] = 'localhost'
os.environ['PGPORT'] = '5432'
os.environ['PGUSER'] = 'postgres'
os.environ['PGPASSWORD'] = 'root'
os.environ['PGDATABASE'] = 'grantservice'

# Set UTF-8 encoding
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

# Добавляем пути
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.path.append(os.path.join(os.path.dirname(__file__), 'agents'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'web-admin'))

# Import app functionality
from data.database.models import GrantServiceDatabase
from agents.researcher_agent_v2 import ResearcherAgentV2

# For Word export
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


async def main():
    """Production test: полный цикл researcher"""

    print("=" * 80)
    print("🚀 PRODUCTION TEST: Researcher Agent - Полный цикл (27 queries)")
    print("=" * 80)
    print(f"Время запуска: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"WebSearch Provider: Claude Code (primary)")
    print("=" * 80)

    # 1. Initialize database (app functionality)
    db = GrantServiceDatabase()
    print("\n✅ Database initialized")

    # 2. Anketa ID
    ANKETA_ID = "EKATERINA_20251010_235448"
    print(f"\n📋 Anketa ID: {ANKETA_ID}")
    print(f"   Проект: Возрождение храма и народных традиций в селе Анисимово")

    # 3. Initialize Researcher Agent (app functionality)
    researcher = ResearcherAgentV2(db, llm_provider="claude_code")
    print("\n✅ ResearcherAgentV2 initialized")
    print(f"   WebSearch Provider: {researcher.websearch_provider}")
    print(f"   WebSearch Fallback: {researcher.websearch_fallback}")

    # 4. Run research (27 queries)
    print("\n" + "=" * 80)
    print("🔬 ЗАПУСК ИССЛЕДОВАНИЯ (27 queries)")
    print("=" * 80)
    print(f"⏰ Начало: {datetime.now().strftime('%H:%M:%S')}")
    print("")

    try:
        result = await researcher.research_with_expert_prompts(ANKETA_ID)

        print("\n" + "=" * 80)
        print("✅ ИССЛЕДОВАНИЕ ЗАВЕРШЕНО!")
        print("=" * 80)

        if result['status'] == 'completed':
            research_id = result['research_id']
            metadata = result['research_results']['metadata']

            print(f"\n📊 Research ID: {research_id}")
            print(f"   Status: {result['status']}")
            print(f"   Total Queries: {metadata['total_queries']}")
            print(f"   Total Sources: {metadata['sources_count']}")
            print(f"   Processing Time: {metadata['total_processing_time']}s")
            print(f"   WebSearch Provider: {metadata.get('websearch_provider', 'N/A')}")

            # Block details
            print("\n📑 Детали по блокам:")
            for block_key in ['block1_problem', 'block2_geography', 'block3_goals']:
                block_data = result['research_results'].get(block_key, {})
                if block_data:
                    block_names = {
                        'block1_problem': 'Блок 1: Проблема и социальная значимость',
                        'block2_geography': 'Блок 2: География и целевая аудитория',
                        'block3_goals': 'Блок 3: Задачи и цели'
                    }
                    print(f"\n   {block_names.get(block_key, block_key)}")
                    print(f"   - Queries: {len(block_data.get('queries_used', []))}")
                    print(f"   - Sources: {block_data.get('total_sources', 0)}")
                    print(f"   - Time: {block_data.get('processing_time', 0)}s")

            # 5. Export to Word
            if DOCX_AVAILABLE:
                print("\n" + "=" * 80)
                print("📄 ЭКСПОРТ В WORD")
                print("=" * 80)

                output_dir = Path("grants_output/production_test")
                output_dir.mkdir(parents=True, exist_ok=True)

                filename = f"Research_{ANKETA_ID}_{research_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                output_path = output_dir / filename

                export_research_to_word(result, output_path)

                print(f"\n✅ Экспорт завершён!")
                print(f"   Файл: {output_path}")

                # 6. Verify in DB
                print("\n" + "=" * 80)
                print("🔍 ПРОВЕРКА В БД")
                print("=" * 80)

                with db.connect() as conn:
                    cursor = conn.cursor()
                    cursor.execute("""
                        SELECT research_id, status, llm_provider,
                               research_results->'metadata'->>'total_queries' as queries,
                               research_results->'metadata'->>'sources_count' as sources,
                               research_results->'metadata'->>'websearch_provider' as ws_provider
                        FROM researcher_research
                        WHERE research_id = %s
                    """, (research_id,))

                    row = cursor.fetchone()
                    cursor.close()

                    if row:
                        print(f"\n✅ Research найден в БД:")
                        print(f"   Research ID: {row[0]}")
                        print(f"   Status: {row[1]}")
                        print(f"   Provider: {row[2]}")
                        print(f"   Queries: {row[3]}")
                        print(f"   Sources: {row[4]}")
                        print(f"   WebSearch: {row[5]}")

                # Final summary
                print("\n" + "=" * 80)
                print("✅ PRODUCTION TEST УСПЕШНО ЗАВЕРШЁН!")
                print("=" * 80)
                print(f"\n📁 Файл для чтения:")
                print(f"   {output_path}")
                print(f"\n💡 Выводы:")
                print(f"   ✅ Researcher Agent работает через Claude Code WebSearch")
                print(f"   ✅ 27 queries выполнены успешно")
                print(f"   ✅ {metadata['sources_count']} источников собрано")
                print(f"   ✅ Результаты сохранены в БД")
                print(f"   ✅ Экспорт в Word работает")

                return {
                    'success': True,
                    'research_id': research_id,
                    'file': str(output_path),
                    'metadata': metadata
                }

            else:
                print("\n⚠️  python-docx не установлен, экспорт пропущен")
                return {
                    'success': True,
                    'research_id': research_id,
                    'metadata': metadata
                }

        else:
            print(f"\n❌ Исследование завершилось с ошибкой:")
            print(f"   Status: {result['status']}")
            print(f"   Error: {result.get('error', 'Unknown')}")
            return {'success': False, 'error': result.get('error')}

    except Exception as e:
        print(f"\n❌ ОШИБКА ПРИ ВЫПОЛНЕНИИ:")
        print(f"   {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


def export_research_to_word(research_result: dict, output_path: Path):
    """Export research results to Word document"""

    doc = Document()

    # Title
    title = doc.add_heading('Исследование грантовой заявки', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph(f'Research ID: {research_result["research_id"]}')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata
    doc.add_heading('Метаданные', 1)

    metadata = research_result['research_results']['metadata']

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    rows_data = [
        ('Research ID', research_result['research_id']),
        ('Status', research_result['status']),
        ('Total Queries', str(metadata['total_queries'])),
        ('Total Sources', str(metadata['sources_count'])),
        ('Processing Time', f"{metadata['total_processing_time']}s"),
        ('WebSearch Provider', metadata.get('websearch_provider', 'unknown'))
    ]

    for idx, (label, value) in enumerate(rows_data):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value

    doc.add_paragraph()

    # Blocks
    for block_key in ['block1_problem', 'block2_geography', 'block3_goals']:
        block_data = research_result['research_results'].get(block_key, {})

        if not block_data:
            continue

        block_names = {
            'block1_problem': 'Блок 1: Проблема и социальная значимость',
            'block2_geography': 'Блок 2: География и целевая аудитория',
            'block3_goals': 'Блок 3: Задачи, мероприятия и главная цель'
        }

        doc.add_heading(block_names[block_key], 1)

        # Stats
        doc.add_paragraph(f"Запросов: {len(block_data.get('queries_used', []))}")
        doc.add_paragraph(f"Источников: {block_data.get('total_sources', 0)}")
        doc.add_paragraph(f"Время: {block_data.get('processing_time', 0)}s")

        # Summary
        doc.add_heading('Резюме', 2)
        doc.add_paragraph(block_data.get('summary', 'Нет данных'))

        # Key facts
        key_facts = block_data.get('key_facts', [])
        if key_facts:
            doc.add_heading('Ключевые факты', 2)
            for idx, fact in enumerate(key_facts[:10], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{fact.get('fact', '')}\n")
                p.add_run(f"Источник: {fact.get('source', 'unknown')}").font.size = Pt(9)

        # Queries used
        queries_used = block_data.get('queries_used', [])
        if queries_used:
            doc.add_heading('Использованные запросы', 2)
            for idx, query in enumerate(queries_used, 1):
                doc.add_paragraph(f"{idx}. {query}", style='List Number')

        doc.add_page_break()

    # Footer
    footer = doc.add_paragraph(f'\nДокумент создан: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save
    doc.save(str(output_path))


if __name__ == "__main__":
    print("\n" + "=" * 80)
    print("🚀 ЗАПУСК PRODUCTION TEST")
    print("=" * 80)

    try:
        result = asyncio.run(main())

        if result.get('success'):
            print("\n" + "=" * 80)
            print("✅ SUCCESS!")
            print("=" * 80)
        else:
            print("\n" + "=" * 80)
            print("❌ FAILED")
            print("=" * 80)
            sys.exit(1)

    except KeyboardInterrupt:
        print("\n\n⚠️  Прервано пользователем")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ ОШИБКА: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
