#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Artifact Exporter - экспорт артефактов грантовой заявки в разные форматы

Поддерживаемые форматы:
- TXT - текстовый файл (человекочитаемый)
- PDF - PDF документ (для печати)
- DOCX - Word документ (для редактирования)
"""

import logging
import io
from typing import Dict, Any, BinaryIO
from datetime import datetime

logger = logging.getLogger(__name__)


class ArtifactExporter:
    """Экспортер артефактов грантовой заявки"""

    def __init__(self, lifecycle_data: Dict[str, Any]):
        """
        Инициализация экспортера

        Args:
            lifecycle_data: Полные данные жизненного цикла от GrantLifecycleManager
        """
        self.data = lifecycle_data
        self.anketa_id = lifecycle_data.get('anketa_id', 'UNKNOWN')
        self.metadata = lifecycle_data.get('metadata', {})
        self.artifacts = lifecycle_data.get('artifacts', {})

    def export_to_txt(self) -> str:
        """
        Экспорт всех артефактов в текстовый формат

        Returns:
            Строка с текстовым представлением всех артефактов
        """
        lines = []

        # Заголовок
        lines.append("=" * 80)
        lines.append(f"ГРАНТОВАЯ ЗАЯВКА")
        lines.append(f"ID: {self.anketa_id}")
        lines.append("=" * 80)
        lines.append("")

        # Метаданные
        lines.append("📋 МЕТАДАННЫЕ")
        lines.append("-" * 80)
        username = self.metadata.get('username', 'Unknown')
        first_name = self.metadata.get('first_name', '')
        last_name = self.metadata.get('last_name', '')
        full_name = f"{first_name} {last_name}".strip() or "Unknown"

        lines.append(f"Пользователь: @{username} ({full_name})")
        lines.append(f"Telegram ID: {self.metadata.get('telegram_id', 'N/A')}")
        lines.append(f"Начало работы: {self.metadata.get('session_started', 'N/A')}")
        lines.append(f"Последнее обновление: {self.metadata.get('session_updated', 'N/A')}")
        lines.append(f"Прогресс: {self.data.get('progress', 0):.0f}%")
        lines.append(f"Текущий этап: {self.data.get('current_stage', 'N/A')}")
        lines.append("")

        # Этап 1: Интервью
        interview = self.artifacts.get('interview', {})
        if interview.get('status') == 'completed':
            lines.append("📝 ЭТАП 1: ИНТЕРВЬЮ (АНКЕТА)")
            lines.append("-" * 80)
            lines.append(f"Количество вопросов: {interview.get('questions_count', 0)}")
            lines.append(f"Завершено: {interview.get('completed_at', 'N/A')}")
            lines.append("")

            for qa in interview.get('data', []):
                lines.append(f"Q{qa.get('question_id', '?')}: {qa.get('question_text', '')}")
                lines.append(f"A: {qa.get('answer', '')}")
                lines.append("")
        else:
            lines.append("📝 ЭТАП 1: ИНТЕРВЬЮ - Не завершен")
            lines.append("")

        # Этап 2: Аудит
        auditor = self.artifacts.get('auditor', {})
        if auditor.get('status') == 'completed':
            lines.append("🔍 ЭТАП 2: АУДИТ ПРОЕКТА")
            lines.append("-" * 80)
            lines.append(f"Оценка качества: {auditor.get('score', 'N/A')}/10")
            lines.append(f"Завершено: {auditor.get('completed_at', 'N/A')}")
            lines.append("")

            if auditor.get('analysis'):
                lines.append("Детальный анализ:")
                lines.append(str(auditor.get('analysis', '')))
                lines.append("")

            if auditor.get('feasibility'):
                lines.append("Оценка реализуемости:")
                lines.append(str(auditor.get('feasibility', '')))
                lines.append("")

            if auditor.get('risks'):
                lines.append("Факторы риска:")
                lines.append(str(auditor.get('risks', '')))
                lines.append("")

            if auditor.get('recommendations'):
                lines.append("Рекомендации:")
                lines.append(str(auditor.get('recommendations', '')))
                lines.append("")
        else:
            lines.append("🔍 ЭТАП 2: АУДИТ - Не завершен")
            lines.append("")

        # Этап 3: Исследование
        researcher = self.artifacts.get('researcher', {})
        if researcher.get('status') == 'completed':
            lines.append("📊 ЭТАП 3: ИССЛЕДОВАНИЕ")
            lines.append("-" * 80)
            lines.append(f"Завершено: {researcher.get('completed_at', 'N/A')}")
            lines.append("")

            if researcher.get('content'):
                lines.append("Основное содержание:")
                lines.append(str(researcher.get('content', '')))
                lines.append("")

            if researcher.get('market'):
                lines.append("Анализ рынка:")
                lines.append(str(researcher.get('market', '')))
                lines.append("")

            if researcher.get('competitors'):
                lines.append("Анализ конкурентов:")
                lines.append(str(researcher.get('competitors', '')))
                lines.append("")

            if researcher.get('sources'):
                lines.append("Источники:")
                lines.append(str(researcher.get('sources', '')))
                lines.append("")
        else:
            lines.append("📊 ЭТАП 3: ИССЛЕДОВАНИЕ - Не завершен")
            lines.append("")

        # Этап 4: Планирование
        planner = self.artifacts.get('planner', {})
        if planner.get('status') == 'completed':
            lines.append("📋 ЭТАП 4: ПЛАНИРОВАНИЕ СТРУКТУРЫ")
            lines.append("-" * 80)
            lines.append(f"Завершено: {planner.get('completed_at', 'N/A')}")
            lines.append("")

            if planner.get('structure'):
                lines.append("Структура заявки:")
                lines.append(str(planner.get('structure', '')))
                lines.append("")

            if planner.get('sections'):
                lines.append("Секции:")
                sections = planner.get('sections', [])
                if isinstance(sections, list):
                    for i, section in enumerate(sections, 1):
                        lines.append(f"{i}. {section}")
                else:
                    lines.append(str(sections))
                lines.append("")

            if planner.get('recommendations'):
                lines.append("Рекомендации:")
                lines.append(str(planner.get('recommendations', '')))
                lines.append("")
        else:
            lines.append("📋 ЭТАП 4: ПЛАНИРОВАНИЕ - Не завершен")
            lines.append("")

        # Этап 5: Финальный грант
        writer = self.artifacts.get('writer', {})
        if writer.get('status') == 'completed':
            lines.append("✍️ ЭТАП 5: ФИНАЛЬНЫЙ ГРАНТ")
            lines.append("=" * 80)
            lines.append(f"ID Гранта: {writer.get('grant_id', 'N/A')}")
            lines.append(f"Название: {writer.get('title', 'N/A')}")
            lines.append(f"Оценка качества: {writer.get('quality_score', 'N/A')}/10")
            lines.append(f"LLM: {writer.get('llm_provider', 'N/A')} ({writer.get('model', 'N/A')})")
            lines.append(f"Завершено: {writer.get('completed_at', 'N/A')}")
            lines.append("=" * 80)
            lines.append("")

            # Секции гранта
            sections = writer.get('sections', [])
            if sections:
                lines.append("СОДЕРЖАНИЕ ГРАНТА ПО СЕКЦИЯМ:")
                lines.append("")
                for section in sections:
                    if isinstance(section, dict):
                        title = section.get('title', 'Без названия')
                        content = section.get('content', '')
                        lines.append(f"## {title}")
                        lines.append("-" * 80)
                        lines.append(content)
                        lines.append("")
                    else:
                        lines.append(str(section))
                        lines.append("")

            # Полный текст гранта
            if writer.get('content'):
                lines.append("")
                lines.append("ПОЛНЫЙ ТЕКСТ ГРАНТА:")
                lines.append("=" * 80)
                lines.append(writer.get('content', ''))
                lines.append("")
        else:
            lines.append("✍️ ЭТАП 5: ФИНАЛЬНЫЙ ГРАНТ - Не завершен")
            lines.append("")

        # Футер
        lines.append("")
        lines.append("=" * 80)
        lines.append(f"Документ создан: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"GrantService - AI-Powered Grant Application System")
        lines.append("=" * 80)

        return "\n".join(lines)

    def export_to_pdf(self) -> bytes:
        """
        Экспорт всех артефактов в PDF формат

        Returns:
            Байты PDF документа
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_LEFT, TA_CENTER

            # Регистрация русского шрифта
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                font_name = 'DejaVuSans'
            except:
                # Fallback на базовый шрифт
                font_name = 'Helvetica'

            # Создание PDF в памяти
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)

            # Стили
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=16,
                alignment=TA_CENTER
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=14
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10
            )

            # Сборка содержимого
            story = []

            # Заголовок
            story.append(Paragraph("ГРАНТОВАЯ ЗАЯВКА", title_style))
            story.append(Paragraph(f"ID: {self.anketa_id}", normal_style))
            story.append(Spacer(1, 0.5*cm))

            # Используем текстовый экспорт как основу
            txt_content = self.export_to_txt()

            # Разбиваем на параграфы и добавляем в PDF
            for line in txt_content.split('\n'):
                if line.strip():
                    if line.startswith('=') or line.startswith('-'):
                        story.append(Spacer(1, 0.3*cm))
                    elif line.startswith('##'):
                        story.append(Paragraph(line.replace('##', ''), heading_style))
                    elif any(line.startswith(emoji) for emoji in ['📋', '📝', '🔍', '📊', '✍️']):
                        story.append(Spacer(1, 0.5*cm))
                        story.append(Paragraph(line, heading_style))
                    else:
                        story.append(Paragraph(line, normal_style))

            # Генерация PDF
            doc.build(story)

            pdf_bytes = buffer.getvalue()
            buffer.close()

            return pdf_bytes

        except ImportError:
            logger.error("reportlab not installed. Install: pip install reportlab")
            # Возвращаем текстовую версию как fallback
            return self.export_to_txt().encode('utf-8')
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            return self.export_to_txt().encode('utf-8')

    def export_to_docx(self) -> bytes:
        """
        Экспорт всех артефактов в DOCX формат

        Returns:
            Байты DOCX документа
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            # Создание документа
            doc = Document()

            # Заголовок
            title = doc.add_heading('ГРАНТОВАЯ ЗАЯВКА', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            subtitle = doc.add_paragraph(f'ID: {self.anketa_id}')
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Метаданные
            doc.add_heading('📋 Метаданные', level=1)

            username = self.metadata.get('username', 'Unknown')
            first_name = self.metadata.get('first_name', '')
            last_name = self.metadata.get('last_name', '')
            full_name = f"{first_name} {last_name}".strip() or "Unknown"

            doc.add_paragraph(f'Пользователь: @{username} ({full_name})')
            doc.add_paragraph(f'Telegram ID: {self.metadata.get("telegram_id", "N/A")}')
            doc.add_paragraph(f'Прогресс: {self.data.get("progress", 0):.0f}%')

            # Этапы
            for stage_name, stage_emoji, stage_key in [
                ('ИНТЕРВЬЮ (АНКЕТА)', '📝', 'interview'),
                ('АУДИТ ПРОЕКТА', '🔍', 'auditor'),
                ('ИССЛЕДОВАНИЕ', '📊', 'researcher'),
                ('ПЛАНИРОВАНИЕ СТРУКТУРЫ', '📋', 'planner'),
                ('ФИНАЛЬНЫЙ ГРАНТ', '✍️', 'writer')
            ]:
                artifact = self.artifacts.get(stage_key, {})

                doc.add_page_break()
                doc.add_heading(f'{stage_emoji} {stage_name}', level=1)

                if artifact.get('status') == 'completed':
                    # Специфичная обработка для каждого этапа
                    if stage_key == 'interview':
                        doc.add_paragraph(f"Вопросов: {artifact.get('questions_count', 0)}")
                        for qa in artifact.get('data', []):
                            q = doc.add_paragraph()
                            q.add_run(f"Q{qa.get('question_id', '?')}: ").bold = True
                            q.add_run(qa.get('question_text', ''))

                            a = doc.add_paragraph()
                            a.add_run('A: ').bold = True
                            a.add_run(qa.get('answer', ''))

                    elif stage_key == 'auditor':
                        doc.add_paragraph(f"Оценка: {artifact.get('score', 'N/A')}/10")
                        if artifact.get('analysis'):
                            doc.add_paragraph(str(artifact.get('analysis', '')))

                    elif stage_key == 'writer':
                        doc.add_paragraph(f"Название: {artifact.get('title', 'N/A')}")
                        doc.add_paragraph(f"Оценка: {artifact.get('quality_score', 'N/A')}/10")

                        # Секции гранта
                        for section in artifact.get('sections', []):
                            if isinstance(section, dict):
                                doc.add_heading(section.get('title', 'Без названия'), level=2)
                                doc.add_paragraph(section.get('content', ''))

                        # Полный текст
                        if artifact.get('content'):
                            doc.add_heading('Полный текст гранта', level=2)
                            doc.add_paragraph(artifact.get('content', ''))

                    else:
                        # Общая обработка для других этапов
                        for key, value in artifact.items():
                            if key not in ['status', 'completed_at'] and value:
                                doc.add_paragraph(f'{key}: {value}')
                else:
                    doc.add_paragraph('Этап не завершен')

            # Футер
            doc.add_page_break()
            footer = doc.add_paragraph(f'Документ создан: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Сохранение в память
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            return docx_bytes

        except ImportError:
            logger.error("python-docx not installed. Install: pip install python-docx")
            # Возвращаем текстовую версию как fallback
            return self.export_to_txt().encode('utf-8')
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return self.export_to_txt().encode('utf-8')


def export_artifact(lifecycle_data: Dict[str, Any], format: str = 'txt') -> bytes:
    """
    Экспортировать артефакты в указанный формат

    Args:
        lifecycle_data: Данные жизненного цикла от GrantLifecycleManager
        format: Формат экспорта ('txt', 'pdf', 'docx')

    Returns:
        Байты документа в указанном формате
    """
    exporter = ArtifactExporter(lifecycle_data)

    if format.lower() == 'pdf':
        return exporter.export_to_pdf()
    elif format.lower() == 'docx':
        return exporter.export_to_docx()
    else:  # txt по умолчанию
        return exporter.export_to_txt().encode('utf-8')
